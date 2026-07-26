# -*- coding: utf-8 -*-
import os, sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from generate_infographics_v2 import build_infographic_s35_v2, build_infographic_s36_v2, build_infographic_s37_v2

FONT_EN = "Arial"
FONT_ZH = "SimHei"  # 中黑體

def set_style_font_and_color(doc, style_name, font_size, bold=False, color_rgb=None, font_en=FONT_EN, font_zh=FONT_ZH):
    styles = doc.styles
    if style_name in styles:
        st = styles[style_name]
        st.font.name = font_en
        st.font.size = font_size
        st.font.bold = bold
        if color_rgb:
            st.font.color.rgb = color_rgb
            
        rPr = st.element.get_or_add_rPr()
        for child in list(rPr):
            if child.tag.endswith('rFonts'):
                rPr.remove(child)
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_en}" w:hAnsi="{font_en}" w:eastAsia="{font_zh}" w:cs="{font_en}"/>')
        rPr.insert(0, rFonts)

def set_run_fonts(run, font_name_en=FONT_EN, font_name_zh=FONT_ZH, font_size=None, bold=None, italic=None, color_rgb=None):
    if font_size is not None:
        run.font.size = font_size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
        
    rPr = run._r.get_or_add_rPr()
    for child in list(rPr):
        if child.tag.endswith('rFonts'):
            rPr.remove(child)
            
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name_en}" w:hAnsi="{font_name_en}" w:eastAsia="{font_name_zh}" w:cs="{font_name_en}"/>')
    rPr.insert(0, rFonts)

def add_references_menu_auto_toc(doc, instruction='TOC \\o "1-3" \\h \\z \\u'):
    sdt_xml = f'''
    <w:sdt {nsdecls("w")}>
      <w:sdtPr>
        <w:docPartObj>
          <w:docPartGallery w:val="Table of Contents"/>
          <w:docPartUnique/>
        </w:docPartObj>
      </w:sdtPr>
      <w:sdtContent>
        <w:p>
          <w:fldSimple w:instr="{instruction.replace('"', '&quot;')}"/>
        </w:p>
      </w:sdtContent>
    </w:sdt>
    '''
    doc._element.body.append(parse_xml(sdt_xml))

def add_footer_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        r_lbl = p.add_run("第 ")
        set_run_fonts(r_lbl, font_size=Pt(9), color_rgb=RGBColor(100, 116, 139))
        
        fldSimple = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>')
        p._p.append(fldSimple)
        
        r_lbl2 = p.add_run(" 頁")
        set_run_fonts(r_lbl2, font_size=Pt(9), color_rgb=RGBColor(100, 116, 139))

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_heading_styled(doc, text, level):
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    style_name = style_map.get(level, 'Heading 1')
    
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(14 if level==1 else (10 if level==2 else 8))
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(text)
    if level == 1:
        set_run_fonts(run, font_size=Pt(20), bold=True, color_rgb=RGBColor(12, 35, 64))
    elif level == 2:
        set_run_fonts(run, font_size=Pt(14.5), bold=True, color_rgb=RGBColor(37, 99, 235))
    elif level == 3:
        set_run_fonts(run, font_size=Pt(12), bold=True, color_rgb=RGBColor(30, 58, 138))
    return p

def add_body_p(doc, text, bold_prefix="", space_after=6):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        set_run_fonts(r_b, font_size=Pt(11), bold=True, color_rgb=RGBColor(12, 35, 64))
        
    r = p.add_run(text)
    set_run_fonts(r, font_size=Pt(11), color_rgb=RGBColor(51, 65, 85))
    return p

def build_references_auto_toc_docx():
    build_infographic_s35_v2()
    build_infographic_s36_v2()
    build_infographic_s37_v2()
    
    doc = Document()
    
    # Configure Built-in Word Styles
    set_style_font_and_color(doc, 'Normal', Pt(11), color_rgb=RGBColor(51, 65, 85))
    set_style_font_and_color(doc, 'Heading 1', Pt(20), bold=True, color_rgb=RGBColor(12, 35, 64))
    set_style_font_and_color(doc, 'Heading 2', Pt(14.5), bold=True, color_rgb=RGBColor(37, 99, 235))
    set_style_font_and_color(doc, 'Heading 3', Pt(12), bold=True, color_rgb=RGBColor(30, 58, 138))

    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Footer Page Number
    add_footer_page_number(doc)

    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    p_cov_top = doc.add_paragraph(style='Normal')
    p_cov_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_top.paragraph_format.space_before = Pt(36)
    
    r_badge = p_cov_top.add_run("🛡️ 國家級主權 AI 國防安全評測與驗證體系\n")
    set_run_fonts(r_badge, font_size=Pt(14), bold=True, color_rgb=RGBColor(37, 99, 235))

    r_title = p_cov_top.add_run("AIEC 國防領域 AI 應用導入專案規劃書")
    set_run_fonts(r_title, font_size=Pt(28), bold=True, color_rgb=RGBColor(12, 35, 64))

    p_cov_sub = doc.add_paragraph(style='Normal')
    p_cov_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_sub.paragraph_format.space_before = Pt(16)
    p_cov_sub.paragraph_format.space_after = Pt(64)
    
    r_sub = p_cov_sub.add_run("—— 非對稱作戰態勢下建構可靠、可信任、可解釋與可當責之主權 AI 戰術評測架構 ——")
    set_run_fonts(r_sub, font_size=Pt(13.5), bold=True, color_rgb=RGBColor(71, 85, 105))

    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("專案名稱", "AIEC 國防領域 AI 應用導入與 24 個月專案規劃書"),
        ("規劃觀點", "國防 AI 應用技術長 (CTO) 暨系統總架構師視角"),
        ("合規標準", "ISO 42001 (AIMS)、CMMC Level 2、DoD CDAO T&E、MITRE ATLAS"),
        ("適用戰術", "非對稱作戰、無人機/船/潛艦/車/機器狗、Cyber Range 對抗演練"),
        ("發行版本與日期", "正式發行版 v1.0  |  2026 年 7 月")
    ]
    for idx, (k, v) in enumerate(meta_data):
        row = table_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(1.8)
        c1.width = Inches(4.5)
        
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(k + "：")
        set_run_fonts(r0, font_size=Pt(10.5), bold=True, color_rgb=RGBColor(37, 99, 235))

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        set_run_fonts(r1, font_size=Pt(10.5), color_rgb=RGBColor(12, 35, 64))

    doc.add_page_break()

    # ----------------------------------------------------
    # WORD REFERENCES MENU AUTOMATIC TABLE OF CONTENTS (w:sdt)
    # ----------------------------------------------------
    add_heading_styled(doc, "📋 目錄 (Table of Contents)", level=1)
    add_references_menu_auto_toc(doc, instruction='TOC \\o "1-3" \\h \\z \\u')

    # Pre-rendered preview for instant reading
    toc_items = [
        ("第一章 前言：國防 AI 應用的戰略急迫性與發展理念", "3"),
        ("    1.1 AI 於國防領域應用的必要性與急迫性", "3"),
        ("    1.2 美國國防部 (DoD) AI First 政策理念與經驗啟示", "4"),
        ("    1.3 台灣面對中國威脅與非對稱作戰 (Asymmetric Warfare) 態勢", "5"),
        ("    1.4 無人化打擊載具應用：無人機、無人船、無人潛艦、無人車與機器狗", "6"),
        ("    1.5 實現「可靠、可信任、可解釋、可當責」之國防 AI", "7"),
        ("    1.6 適合台灣本土環境落地方案：為什麼要做、如何做、怎麼樣做", "8"),
        ("第二章 專案執行範圍與三大核心區塊架構", "9"),
        ("    2.1 第一部分：治理與標準 (ISO 42001, SHIELD, RoE, CMMC L2)", "9"),
        ("    2.2 第二部分：矩陣與架構 (T&E 4層階梯, JATIC 7構面, Q1~Q15 量化門檻)", "11"),
        ("    2.3 第三部分：應用系統與驗測 SOP (A~F 類系統與無人載具驗測)", "14"),
        ("第三章 階段性任務與 24 個月時程規劃", "17"),
        ("    3.1 四大階段任務與發展里程碑 (Milestones)", "17"),
        ("    3.2 24 個月專案時程對照表 (Gantt Schedule)", "19"),
        ("第四章 資源配置與專案風險管理", "20"),
        ("    4.1 軟硬體、民雄算力與人力資源配置表", "20"),
        ("    4.2 專案風險評估與因應對策 (Risk Matrix)", "21"),
        ("第五章 預期效益與戰略成果", "22")
    ]
    for title_text, page_num in toc_items:
        p_t = doc.add_paragraph(style='Normal')
        p_t.paragraph_format.space_after = Pt(2)
        r_t = p_t.add_run(f"{title_text}  .........................................................................................................................  {page_num}")
        set_run_fonts(r_t, font_size=Pt(10), bold=True if not title_text.startswith(" ") else False, color_rgb=RGBColor(12, 35, 64))

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ----------------------------------------------------
    # LIST OF FIGURES (圖目錄 - 原生自動生成)
    # ----------------------------------------------------
    add_heading_styled(doc, "🖼️ 圖目錄 (List of Figures)", level=2)
    add_references_menu_auto_toc(doc, instruction='TOC \\t "Heading 3,1" \\h \\z \\u')

    fig_items = [
        ("圖 2-1：AIEC 國防 AI 評測五大階段標準流程圖 (Nano Banana pro 生成)", "10"),
        ("圖 2-2：國防 AI 安全保密審計 3D 縱深防禦矩陣圖 (Nano Banana pro 生成)", "13"),
        ("圖 2-3：民雄院區國家級主權 AI 四層算力堆疊架構圖 (Nano Banana pro 生成)", "16")
    ]
    for title_text, page_num in fig_items:
        p_f = doc.add_paragraph(style='Normal')
        p_f.paragraph_format.space_after = Pt(2)
        r_f = p_f.add_run(f"{title_text}  ...................................................................................................  {page_num}")
        set_run_fonts(r_f, font_size=Pt(10), color_rgb=RGBColor(37, 99, 235))

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ----------------------------------------------------
    # LIST OF TABLES (表目錄 - 原生自動生成)
    # ----------------------------------------------------
    add_heading_styled(doc, "📊 表目錄 (List of Tables)", level=2)
    add_references_menu_auto_toc(doc, instruction='TOC \\t "Heading 3,1" \\h \\z \\u')

    tbl_items = [
        ("表 1-1：傳統軍事戰術 vs. 非對稱 AI 戰術比較表", "5"),
        ("表 2-1：ISO 42001 (AIMS) 通用標準 vs. AIEC 國防權責分工比較表", "9"),
        ("表 2-2：交戰規則 (RoE) 人機授權三階權能邊界比較表 (HITL / HOTL / HOOTL)", "10"),
        ("表 2-3：AIEC 15 項量化評測指標 (Q1~Q15) 與合格門檻對照總表", "12"),
        ("表 3-1：AIEC 國防 AI 導入專案 24 個月執行時程規劃表", "19"),
        ("表 4-1：專案人力、主權算力與軟硬體資源配置表", "20"),
        ("表 4-2：國防 AI 導入主要風險評估與處置矩陣表", "21")
    ]
    for title_text, page_num in tbl_items:
        p_tb = doc.add_paragraph(style='Normal')
        p_tb.paragraph_format.space_after = Pt(2)
        r_tb = p_tb.add_run(f"{title_text}  ...................................................................................................  {page_num}")
        set_run_fonts(r_tb, font_size=Pt(10), color_rgb=RGBColor(30, 58, 138))

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 1: INTRODUCTION
    # ----------------------------------------------------
    add_heading_styled(doc, "第一章 前言：國防 AI 應用的戰略急迫性與發展理念", level=1)
    
    add_heading_styled(doc, "1.1 AI 於國防領域應用的必要性與急迫性", level=2)
    add_body_p(doc, "當前全球軍事科技正面臨百年未有之大變局。人工智慧 (AI) 已非僅是輔助軍隊行政或後勤維運的資訊工具，而是直接主導數位化戰場「擊殺鏈」(Kill Chain) 閉合速度、情報感知（ISR）精度與指揮管制 (C2) 決策品質的核心戰術作戰要素。在高強度對抗之現代化戰爭中，戰術反應時間已被壓縮至秒級甚至毫秒級，人腦處理海量多源傳感器（雷達、光學籌載、電戰訊號）的極限已成為戰術瓶頸，唯有導入 AI 技術，方能實現預警即反應、目標即摧毀之戰術目標。")
    
    add_heading_styled(doc, "1.2 美國國防部 (DoD) AI First 政策理念與經驗啟示", level=2)
    add_body_p(doc, "美國國防部 (DoD) 為確保全球軍事與科技領導地位，已明確頒布「AI First」政策理念，將 AI 提升至國家安全戰略之核心柱石。DoD 設立數位與人工智慧長辦公室 (CDAO)，全面推動 Joint All-Domain Command and Control (JADC2) 全領域聯合指管架構，並透過數據、分析與 AI 進步指導方針，強調將 AI 技術深度嵌入各軍種作戰裝備與戰術決策程序中。美軍的經驗明確示範：AI 導入國防不能採取傳統資訊系統採購邏輯，必須建立嚴格的測試與評估 (T&E) 體系、資料與模型溯源 (Data & Model Provenance) 機制，以及可信賴的負責任 AI (RAI) 驗證流程。")

    add_heading_styled(doc, "1.3 台灣面對中國威脅與非對稱作戰 (Asymmetric Warfare) 態勢", level=2)
    add_body_p(doc, "面對中國解放軍在海陸空天電各領域的強大軍事威脅與數量優勢，台灣在軍事防禦上必須堅定貫徹「非對稱作戰」(Asymmetric Warfare) 戰略。在兵力規模與傳統重型武器數量對比懸殊的態勢下，台灣無法也不應與敵進行消耗性常規戰。發揮「以小搏大、以智勝強」之非對稱優勢，核心關鍵即在於運用高密度、低成本、具備自主運算能力之 AI 無人化系統與地端防衛陣地，對敵形成多區域、多層次之區域拒止與封鎖 (A2/AD)。")

    # Table 1-1 Title (Styled as Heading 3 for Auto-TOC)
    add_heading_styled(doc, "表 1-1：傳統軍事戰術 vs. 非對稱 AI 戰術比較表", level=3)

    t1 = doc.add_table(rows=5, cols=3)
    set_table_borders(t1)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["分析維度", "傳統常規軍事戰術", "非對稱 AI 戰術 (AIEC 標竿)"]
    for j, h in enumerate(headers):
        cell = t1.rows[0].cells[j]
        set_cell_background(cell, "0C2340")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_fonts(r, font_size=Pt(10), bold=True, color_rgb=RGBColor(255, 255, 255))

    rows_t1 = [
        ("戰術思維與戰力表現", "高度依賴常規大型載具與人力，戰力線性成長", "兵力倍增器 (Force Multiplier)，彈性集群飽和打擊"),
        ("擊殺鏈 (Kill Chain) 時延", "人工判讀多源傳感器，反應時延數分鐘至數小時", "AI 邊緣即時目標辨識與自動打擊提案，壓縮至秒級"),
        ("人力風險與成本結構", "人員傷亡風險高，高單價重型武器維護極為昂貴", "無人化載具零傷亡，低成本可消耗性裝備具高防衛效益"),
        ("資安與保密防護能力", "傳統實體金鑰管理，設備俘獲易致演算法洩漏", "網狀防篡改、邊緣模型 100ms 緊急自毀與零降密洩漏")
    ]
    for i, row_data in enumerate(rows_t1):
        row = t1.rows[i+1]
        bg_hex = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_fonts(r, font_size=Pt(9.5), color_rgb=RGBColor(30, 41, 59))

    add_heading_styled(doc, "1.4 無人化打擊載具應用：無人機、無人船、無人潛艦、無人車與機器狗", level=2)
    add_body_p(doc, "在非對稱作戰體系中，五大無人化作戰載具構成立體防衛網的核心。本專案將全面導入 AI 評測 SOP 於下列載具系統：")
    add_body_p(doc, "無人機 (UAV) 具備飛行航路自動規劃、邊緣視覺 (YOLO/Mobile SAM) 目標偵測、紅軍對抗擾動防禦與抗天候降質能力；於無 GPS 網狀環境下實施多機蜂群協同。", bold_prefix="1. 無人機 (UAV) 蜂群戰術系統：")
    add_body_p(doc, "無人船 (USV) 與無人潛艦 (UUV) 部署於台海周邊海域與潛艦航道，具備水面光學/雷達識別、水聲聲納特徵AI識別與自主巡航能力；嚴格遵循交戰規則 (RoE) 劃分。", bold_prefix="2. 無人船 (USV) 與無人潛艦 (UUV) 海上封鎖系統：")
    add_body_p(doc, "無人車 (UGV) 與機器狗 (Robotic Dogs) 用於灘岸防衛、巷戰偵察與設施巡邏。具備複雜地形動態避障、熱成像目標歸因 (Point Game) 與遭敵俘獲時 <100ms 硬體模型自毀零化保護。", bold_prefix="3. 無人車 (UGV) 與戰術機器狗 (Robotic Dogs) 陸戰系統：")

    add_heading_styled(doc, "1.5 實現「可靠、可信任、可解釋、可當責」之國防 AI", level=2)
    add_body_p(doc, "軍事作戰容錯率為零。AI 系統若出現誤判、對抗擾動崩潰或黑盒子無解釋決策，將引發災難性戰術後果。本專案導入 AIEC 評測體系，旨在達成四大核心目標：")
    add_body_p(doc, "系統遭受敵方對抗貼片、電戰雜訊或惡劣天候時，仍維持 90% 以上識別精確度 (Q1, Q2)。", bold_prefix="• 可靠 (Reliable)：")
    add_body_p(doc, "指揮官了解模型信心分數 (Confidence Score)，期望校準誤差 ECE ≤ 0.05，防止盲目過度信任 (Q5)。", bold_prefix="• 可信任 (Trustworthy)：")
    add_body_p(doc, "白箱評測提供 XAI 特徵歸因熱力圖與 Point Game 得分 (≥ 0.85)，使指揮官能明瞭 AI 提案邏輯 (Q7)。", bold_prefix="• 可解釋 (Explainable)：")
    add_body_p(doc, "落實 Data & Model Provenance 完整溯源與高保真日誌保存，明確劃分 RoE 人機授權責任 (HITL/HOTL)。", bold_prefix="• 可當責 (Accountable)：")

    add_heading_styled(doc, "1.6 適合台灣本土環境落地方案：為什麼要做、如何做、怎麼樣做", level=2)
    add_body_p(doc, "為什麼我們必須做？因為面對敵方龐大威脅與資安滲透，台灣絕不能依賴未經確效之商業 AI 或國外雲端 API。我們如何做？由國家中山科學研究院 (NCSIST) 主導主權算力與 AIEC 評測中心，結合國家級民雄院區 100% 地端實體隔離 (Air-Gapped) 主權算力，建立通用國際標準 (ISO 42001) 與國防放行閘門 (TRL) 之雙支柱機制。怎麼樣做最適合台灣環境？實施聯邦學習 (Federated Learning)「模型移動，資料不動」，保護各雷達站與陣地敏感數據，並於戰術邊緣部署硬體防篡改與 <100ms 物理零化自毀，達成極致保密與戰術自主。")

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 2: PROJECT SCOPE & THREE PILLARS
    # ----------------------------------------------------
    add_heading_styled(doc, "第二章 專案執行範圍與三大核心區塊架構", level=1)
    
    add_heading_styled(doc, "2.1 第一部分：治理與標準 (Governance & Standards)", level=2)
    add_body_p(doc, "第一部分建立國防 AI 應用的制度安全底座。融合國際標準 ISO 42001 (AIMS) 與國家級 NCSIST AIEC 總體藍圖，劃分 ISO 通用方法論與國防 AIEC TRL 戰術放行閘門。導入 SHIELD 六大治理生命週期 (Set, Hone, Improve, Evaluate, Log, Detect)，對齊 CMMC Level 2 資安標準，並針對 RAG 系統建立向量 RBAC 標籤動態遮罩，達到零降密洩漏。同時明訂 RoE 交戰規則三階人機授權邊界。")

    # Table 2-1 Title (Styled as Heading 3 for Auto-TOC)
    add_heading_styled(doc, "表 2-1：ISO 42001 (AIMS) 通用標準 vs. AIEC 國防權責分工比較表", level=3)

    t21 = doc.add_table(rows=4, cols=3)
    set_table_borders(t21)
    t21.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["控制維度", "ISO 42001 (AIMS) 國際標準", "AIEC 國防領域特化審查閘門"]):
        cell = t21.rows[0].cells[j]
        set_cell_background(cell, "0C2340")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_fonts(r, font_size=Pt(10), bold=True, color_rgb=RGBColor(255, 255, 255))

    rows_t21 = [
        ("管理範疇與目標", "企業民用 AI 風險評估與品質管理體系 (AIIA)", "武器系統、C2 指管與無人載具 TRL 戰術放行審查"),
        ("對抗防禦與安全", "Annex A.7 規範防範資料污染與基本資安控制", "MITRE ATLAS 16大戰術演練、Cyber Range 與 <100ms 硬體自毀"),
        ("透明度與當責", "Clause 8.4 要求提供基本 AI 說明與利害關係人說明", "Point Game 得分 ≥ 0.85 白箱歸因熱力圖、RoE 授權與 Data Provenance")
    ]
    for i, row_data in enumerate(rows_t21):
        row = t21.rows[i+1]
        bg_hex = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_fonts(r, font_size=Pt(9.5), color_rgb=RGBColor(30, 41, 59))

    # Table 2-2 Title (Styled as Heading 3 for Auto-TOC)
    add_heading_styled(doc, "表 2-2：交戰規則 (RoE) 人機授權三階權能邊界比較表 (HITL / HOTL / HOOTL)", level=3)

    t22 = doc.add_table(rows=4, cols=4)
    set_table_borders(t22)
    t22.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["授權階層 (RoE Tier)", "人類角色 (Human Role)", "AI 系統權能範圍", "適用國防戰術情境"]):
        cell = t22.rows[0].cells[j]
        set_cell_background(cell, "0C2340")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_fonts(r, font_size=Pt(9.5), bold=True, color_rgb=RGBColor(255, 255, 255))

    rows_t22 = [
        ("HITL (人在紐中)", "最終打擊與開火決策者", "僅提供目標自動識別與 COA 行動方案建議，無發射權限", "飛彈陣地、致傷性武器打擊、高機密情報處置"),
        ("HOTL (人在紐上)", "實時監控者與斷路接管者", "自主執行尋標與防空追蹤，人類具備微秒級強制中斷權", "無人機蜂群對抗干擾、海面無人船自動攔截與巡航"),
        ("HOOTL (完全自主)", "系統範疇設定者", "全自主執行戰術巡航、動態避障與電子偵察", "無人車區域巡邏、戰術機器狗地形探勘、無密級偵察")
    ]
    for i, row_data in enumerate(rows_t22):
        row = t22.rows[i+1]
        bg_hex = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_fonts(r, font_size=Pt(9.0), color_rgb=RGBColor(30, 41, 59))

    # Figure 2-1 Title & Image (Styled as Heading 3 for Auto-TOC)
    p_fig21 = doc.add_paragraph(style='Normal')
    p_fig21.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig21.paragraph_format.space_before = Pt(12)
    p_fig21.paragraph_format.space_after = Pt(4)
    run_fig21_img = p_fig21.add_run()
    run_fig21_img.add_picture(os.path.join('infographics', 'info_pipeline_s35.png'), width=Inches(6.2))
    
    add_heading_styled(doc, "圖 2-1：AIEC 國防 AI 評測五大階段標準流程圖 (Nano Banana pro 生成)", level=3)

    add_heading_styled(doc, "2.2 第二部分：矩陣與架構 (Matrix & Architecture)", level=2)
    add_body_p(doc, "第二部分建立技術評測方法論與量化指標。導入 DoD CDAO Level 1~4 T&E 階梯、JATIC 7 大共通構面（穩健性、韌性、可解釋性、勝任力、公平性、信任校準、漂移監控），並定義 6 大評測方法論（黑箱、白箱、基準測試、紅軍演練、專家評估、持續監控）。建立包含 Q1~Q15 15 項量化指標之 PASS 門檻矩陣，並整合 Security (ATLAS 對抗防禦與自毀)、Confidentiality (地端實體隔離與聯邦學習) 與 Auditing (Data/Model Provenance 溯源與 XAI) 之三維縱深聯防矩陣。")

    # Table 2-3 Title (Styled as Heading 3 for Auto-TOC)
    add_heading_styled(doc, "表 2-3：AIEC 15 項量化評測指標 (Q1~Q15) 與合格門檻對照總表", level=3)

    t23 = doc.add_table(rows=16, cols=4)
    set_table_borders(t23)
    t23.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["編號", "指標名稱 (英文縮寫)", "量化合格判定門檻 (PASS Threshold)", "代表性驗測工具鏈"]):
        cell = t23.rows[0].cells[j]
        set_cell_background(cell, "0C2340")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_fonts(r, font_size=Pt(9.5), bold=True, color_rgb=RGBColor(255, 255, 255))

    q_data = [
        ("Q1", "對抗韌性維持率 (Adversarial Robustness)", "Acc_adv / Acc_clean ≥ 0.90 (ε ≤ 0.05)", "IBM ART 360, HEART, FGSM, PGD"),
        ("Q2", "自然穩健性衰減 (Natural Robustness)", "ΔmAP = (mAP_clean - mAP_noise) / mAP_clean ≤ 0.10", "NRTK Toolkit, ImageNet-C"),
        ("Q3", "任務完成率 (Mission Success Rate)", "MSR = Success Count / N ≥ 0.95 (N=100)", "VBS 4, EADSIM, LVC 平行戰場"),
        ("Q4", "可中止性與失效安全 (Abortability)", "τ_abort ≤ 100ms  ∧  Fail-Safe Rate = 100%", "ToAST, HITL 硬體物理斷路器"),
        ("Q5", "期望校準誤差 (Trust Calibration)", "ECE ≤ 0.05  ∧  R_overreliance ≤ 0.05", "HMT Evaluation Suite, ECE Calculator"),
        ("Q6", "認知負荷優化率 (Cognitive Load)", "ΔTLX ≥ 0.30  ∧  Δt_decision ≤ 2.0s", "NASA-TLX 量表, EEG 腦電儀, 眼動儀"),
        ("Q7", "顯著性歸因得分 (Explainability)", "Point Game Score ≥ 0.85 (85%)", "XAITK, SHAP, LIME, Grad-CAM"),
        ("Q8", "提示越獄阻絕率 (Jailbreak Defense)", "R_jailbreak_def ≥ 0.99 (99%)", "garak LLM Scanner, NeMo Guardrails"),
        ("Q9", "幻覺發生率與忠實度 (Hallucination)", "Faithfulness ≥ 0.95  ∧  Hallucination ≤ 0.02", "RAGAS, TruLens Triad, PromptBench"),
        ("Q10", "RAG 檢索精確與歸屬 (Context Precision)", "Context Precision ≥ 0.90  ∧  Attribution ≥ 0.98", "RAGAS Assessment Suite, TruLens"),
        ("Q11", "Agent 越權調用率 (Tool Misuse)", "R_unauth_API = 0% (零越權呼叫事故)", "AgentBench, Open Policy Agent (OPA)"),
        ("Q12", "概念漂移召回率 (Concept Drift)", "Drift Recall ≥ 0.95  ∧  t_alarm ≤ 5min", "PyOD, Alibi Detect, Evidently AI"),
        ("Q13", "不確定性覆蓋率 (Uncertainty Quant)", "OOD 方差覆蓋率 Coverage ≥ 95%", "MC-Dropout, Deep Ensembles, UQ"),
        ("Q14", "防降密洩漏率 (Declassification Leak)", "R_declass_leak = 0% (零降密洩漏風險)", "Milvus 向量 RBAC 標籤動態遮罩"),
        ("Q15", "系統軌跡可追溯性 (Traceability)", "Log Coverage = 100%  ∧  t_repro ≤ 10min", "OpenTelemetry, SPIFFE/SPIRE 日誌")
    ]
    for i, row_data in enumerate(q_data):
        row = t23.rows[i+1]
        bg_hex = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            if j in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_fonts(r, font_size=Pt(8.5), color_rgb=RGBColor(30, 41, 59))

    # Figure 2-2 Title & Image (Styled as Heading 3 for Auto-TOC)
    p_fig22 = doc.add_paragraph(style='Normal')
    p_fig22.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig22.paragraph_format.space_before = Pt(12)
    p_fig22.paragraph_format.space_after = Pt(4)
    run_fig22_img = p_fig22.add_run()
    run_fig22_img.add_picture(os.path.join('infographics', 'info_3d_matrix_s36.png'), width=Inches(6.2))
    
    add_heading_styled(doc, "圖 2-2：國防 AI 安全保密審計 3D 縱深防禦矩陣圖 (Nano Banana pro 生成)", level=3)

    add_heading_styled(doc, "2.3 第三部分：應用系統與驗測 SOP (Application Systems & SOPs)", level=2)
    add_body_p(doc, "第三部分將評測規範落地至六大類 AI 應用系統與前線無人載具：A類電腦視覺 (YOLO/SAM/NRTK)、B類生成式AI (garak/NeMo)、C類RAG知識庫 (RAGAS/RBAC)、D類AI Agent (SPIFFE/OPA)、E類自主系統與HMT (RoE/VBS4) 及 F類預測分析 (PyOD/UQ)。針對無人機蜂群、無人船、無人潛艦、無人車與機器狗，實施網狀防篡改與 <100ms 硬體自毀評測。")

    # Figure 2-3 Title & Image (Styled as Heading 3 for Auto-TOC)
    p_fig23 = doc.add_paragraph(style='Normal')
    p_fig23.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig23.paragraph_format.space_before = Pt(12)
    p_fig23.paragraph_format.space_after = Pt(4)
    run_fig23_img = p_fig23.add_run()
    run_fig23_img.add_picture(os.path.join('infographics', 'info_4tier_compute_s37.png'), width=Inches(6.2))
    
    add_heading_styled(doc, "圖 2-3：民雄院區國家級主權 AI 四層算力堆疊架構圖 (Nano Banana pro 生成)", level=3)

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 3: PHASED TASKS & SCHEDULE
    # ----------------------------------------------------
    add_heading_styled(doc, "第三章 階段性任務與 24 個月時程規劃", level=1)
    
    add_heading_styled(doc, "3.1 四大階段任務與發展里程碑 (Milestones)", level=2)
    add_body_p(doc, "本專案規劃為期 24 個月之推動時程，劃分為四大循序漸進階段：")
    add_body_p(doc, "完成 ISO 42001 AIMS 通用框架調校、頒布國防 AIEC 評測手冊與 Q1~Q15 量化合格門檻，完成民雄院區 Tier 1~3 地端實體隔離算力硬體建置。", bold_prefix="第一階段：治理標準與主權算力基建整備 (第 1 - 6 個月)：")
    add_body_p(doc, "部署 garak, IBM ART 360, NRTK, RAGAS 自動化評測工具鏈；建置 Cyber Range AI 對抗演練沙盒，完成首批 A、B 類系統紅軍對抗演練。", bold_prefix="第二階段：沙盒演練與自動化工具鏈部署 (第 7 - 12 個月)：")
    add_body_p(doc, "結合 VBS 4 / EADSIM LVC 平行戰場，實施無人機蜂群、無人船、無人潛艦、無人車與機器狗實戰壓力測試；驗證網狀防篡改與 <100ms 模型自毀。", bold_prefix="第三階段：無人載具與多系統實戰壓力驗證 (第 13 - 18 個月)：")
    add_body_p(doc, "全面聯網運作 MLOps 自動化驗測工廠，實現確效紀錄自動轉換為 CMMC L2 控制項與 ISO 42001 審計日誌，核發主權 TRL 戰術放行證書。", bold_prefix="第四階段：全域營運與主權 TRL 放行部署 (第 19 - 24 個月)：")

    add_heading_styled(doc, "3.2 24 個月專案時程對照表 (Gantt Schedule)", level=2)

    # Table 3-1 Title (Styled as Heading 3 for Auto-TOC)
    add_heading_styled(doc, "表 3-1：AIEC 國防 AI 導入專案 24 個月執行時程規劃表", level=3)

    t31 = doc.add_table(rows=9, cols=5)
    set_table_borders(t31)
    t31.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["專案階段", "核心任務工作項目", "時程 (月份)", "負責單位", "階段里程碑 (Milestones)"]):
        cell = t31.rows[0].cells[j]
        set_cell_background(cell, "0C2340")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_fonts(r, font_size=Pt(9.5), bold=True, color_rgb=RGBColor(255, 255, 255))

    sched_data = [
        ("一階段", "ISO 42001 與 AIEC 治理標準擬定", "M1 - M3", "AIEC 治理小組", "發布《國防 AI 評測與 RoE SOP 手冊》"),
        ("一階段", "民雄院區四層算力與地端隔離整備", "M4 - M6", "主權基建小組", "完成 70B/8x7B/7B 機房與 Air-Gap 通電"),
        ("二階段", "garak, ART 360, RAGAS 工具鏈部署", "M7 - M9", "評測工具小組", "上線 MLOps 自動化驗測工具包"),
        ("二階段", "Cyber Range AI 對抗演練沙盒測試", "M10 - M12", "紅隊對抗小組", "完成 10,000 筆越獄與對抗貼片測試"),
        ("三階段", "無人機/船/潛艦/車/機器狗 LVC 測試", "M13 - M15", "無人戰術小組", "完成 100 次蒙地卡羅戰術模擬 (MSR≥0.95)"),
        ("三階段", "<100ms 模型硬體緊急自毀測試", "M16 - M18", "資安防護小組", "驗證物理零化與 Flash/RAM 衝刷率 100%"),
        ("四階段", "CMMC L2 與 ISO 42001 自動化審計稽核", "M19 - M21", "資安合規小組", "通過 CMMC L2 控制項審計對接"),
        ("四階段", "主權 TRL 證書核發與戰術陣地部署放行", "M22 - M24", "AIEC 評測中樞", "核發首批武器系統 TRL 技術放行證書")
    ]
    for i, row_data in enumerate(sched_data):
        row = t31.rows[i+1]
        bg_hex = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p = cell.paragraphs[0]
            if j in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_fonts(r, font_size=Pt(8.5), color_rgb=RGBColor(30, 41, 59))

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 4: RESOURCES & RISK MANAGEMENT
    # ----------------------------------------------------
    add_heading_styled(doc, "第四章 資源配置與專案風險管理", level=1)
    
    add_heading_styled(doc, "4.1 軟硬體、民雄算力與人力資源配置表", level=2)
    
    # Table 4-1 Title (Styled as Heading 3 for Auto-TOC)
    add_heading_styled(doc, "表 4-1：專案人力、主權算力與軟硬體資源配置表", level=3)

    t41 = doc.add_table(rows=5, cols=3)
    set_table_borders(t41)
    t41.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["資源類別", "配置規格與內容", "戰術與評測效益"]):
        cell = t41.rows[0].cells[j]
        set_cell_background(cell, "0C2340")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_fonts(r, font_size=Pt(9.5), bold=True, color_rgb=RGBColor(255, 255, 255))

    res_data = [
        ("主權算力硬體", "民雄院區 100% 地端實體隔離 (Air-Gap) GPU 伺服器叢集", "支援 Tier 1 70B 模型與 Tier 2 8x7B MoE 並行高併發推論"),
        ("評測軟體工具鏈", "garak, IBM ART 360, NRTK, XAITK, RAGAS, AgentBench, PyOD", "覆蓋黑箱、白箱、紅軍演練、XAI 與 15 項量化指標自動評測"),
        ("模擬驗證設施", "VBS 4, EADSIM, ToAST, Cyber Range 平行戰場演練沙盒", "支援 100 次蒙地卡羅模擬、RoE 人機協同與 <100ms 自毀測試"),
        ("專家專業團隊", "國防 AI 技術長暨總架構師、資安 ISO 42001/CMMC 稽核員、紅藍隊專家", "主導全案 SOP 擬定、TRL 審查、演算法後門稽核與責任劃分")
    ]
    for i, row_data in enumerate(res_data):
        row = t41.rows[i+1]
        bg_hex = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_fonts(r, font_size=Pt(9.0), color_rgb=RGBColor(30, 41, 59))

    add_heading_styled(doc, "4.2 專案風險評估與因應對策 (Risk Matrix)", level=2)

    # Table 4-2 Title (Styled as Heading 3 for Auto-TOC)
    add_heading_styled(doc, "表 4-2：國防 AI 導入主要風險評估與處置矩陣表", level=3)

    t42 = doc.add_table(rows=5, cols=4)
    set_table_borders(t42)
    t42.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["風險項目", "風險等級", "可能影響分析", "預防與緩解處置對策"]):
        cell = t42.rows[0].cells[j]
        set_cell_background(cell, "0C2340")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_fonts(r, font_size=Pt(9.5), bold=True, color_rgb=RGBColor(255, 255, 255))

    risk_data = [
        ("對抗貼片與數據污染", "高 (High)", "目標偵測誤判，無人機導彈攻擊無效目標", "部署 IBM ART 360 預處理擾動，執行 Q1 對抗韌性門檻 (≥0.90)"),
        ("戰術邊緣裝備遭俘獲", "極高 (Critical)", "敵方逆向工程獲取權重與保密演算法", "軍規 Enclave 晶片與 <100ms Flash/RAM 雜訊零化緊急自毀"),
        ("敏感研發數據洩漏", "高 (High)", "RAG 檢索輸出不慎暴露機密降密等級", "Milvus 向量 RBAC 標籤動態遮罩，合規 CMMC L2 零降密洩漏"),
        ("操作員過度依賴 AI", "中 (Medium)", "盲目採納高信心錯誤提案引發戰術失誤", "評測 ECE ≤ 0.05 信心校準，嚴格劃分 RoE 人力最終授權")
    ]
    for i, row_data in enumerate(risk_data):
        row = t42.rows[i+1]
        bg_hex = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            if j in [0, 1]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_fonts(r, font_size=Pt(9.0), color_rgb=RGBColor(30, 41, 59))

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER 5: EXPECTED BENEFITS & CONCLUSION
    # ----------------------------------------------------
    add_heading_styled(doc, "第五章 預期效益與戰略成果", level=1)
    
    add_body_p(doc, "本專案執行完畢後，將為國防防衛體系與國內國防產業帶來三大層面之關鍵戰略成果：")
    
    add_heading_styled(doc, "1. 戰術作戰層面：達成極致非對稱打擊效益", level=2)
    add_body_p(doc, "實現無人機蜂群、無人船、無人潛艦、無人車與機器狗之跨域立體協同打擊能力。擊殺鏈反應時間由傳統數分鐘壓縮至秒級；端到端任務完成率 (MSR) 達 95% 以上，顯著提升台海防衛之區域拒止能力。")

    add_heading_styled(doc, "2. 資安與保密層面：建構零漏洞主權防衛屏障", level=2)
    add_body_p(doc, "達成 100% 地端實體隔離 (Air-Gap) 主權算力營運，落實聯邦學習「模型移動，資料不動」；邊緣硬體自毀回應時間控制於 100ms 內；研發數據採 Milvus 向量 RBAC 動態遮罩，合規 CMMC Level 2，達到零降密洩漏風險。")

    add_heading_styled(doc, "3. 國家與產業層面：樹立自主 AIEC 評測認證標竿", level=2)
    add_body_p(doc, "整合 ISO 42001 (AIMS) 與國家級 NCSIST AIEC 總體藍圖，建立國內第一個權威國防 AI 測試與評估 (T&E) 中心，核發技術成熟度 (TRL) 戰術放行證書。不僅加速國內軍工產業 AI 升級，更確保國軍武器裝備具備國際頂級之可靠、可信任、可解釋與可當責品質。")

    # Save local and desktop
    out_local = r'c:\Users\administartor\Downloads\AIEC\AIEC國防領域AI應用導入專案規劃書.docx'
    out_desktop = r'C:\Users\administartor\Desktop\AIEC國防領域AI應用導入專案規劃書.docx'
    
    doc.save(out_local)
    doc.save(out_desktop)
    print(f"Successfully generated Word proposal document with References menu auto-TOC structure & Heading 3 Figure/Table titles at:\n - {out_local}\n - {out_desktop}")

if __name__ == '__main__':
    build_references_auto_toc_docx()
